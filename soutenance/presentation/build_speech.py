# -*- coding: utf-8 -*-
"""Script de soutenance CuraMedical (Word), 3 intervenants génériques, ~13 min.
Vocabulaire simple, ton professionnel et charismatique."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

GREEN = RGBColor(0x2A, 0x9B, 0x69)
DARK  = RGBColor(0x0F, 0x17, 0x2A)
GREY  = RGBColor(0x64, 0x74, 0x8B)
FONT  = "Calibri"

doc = Document()
st = doc.styles['Normal']; st.font.name = FONT; st.font.size = Pt(11); st.font.color.rgb = DARK
for s in doc.sections:
    s.top_margin = Inches(0.8); s.bottom_margin = Inches(0.8)
    s.left_margin = Inches(0.9); s.right_margin = Inches(0.9)

def para(text="", size=11, color=DARK, bold=False, italic=False,
         align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=6, ls=1.15):
    p = doc.add_paragraph(); p.alignment = align
    pf = p.paragraph_format; pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = ls
    if text:
        r = p.add_run(text); r.font.size = Pt(size); r.font.bold = bold
        r.font.italic = italic; r.font.color.rgb = color; r.font.name = FONT
    return p

def runs(p, parts):
    for t, o in parts:
        r = p.add_run(t); r.font.size = Pt(o.get('size', 11)); r.font.bold = o.get('bold', False)
        r.font.italic = o.get('italic', False); r.font.color.rgb = o.get('color', DARK); r.font.name = FONT
    return p

para("Script de soutenance", 26, GREEN, bold=True, after=0)
para("CuraMedical — Plateforme intelligente de gestion de cabinet médical", 13, GREY, after=2)
para("ISGA · 2CI-ISI · Groupe 2 · Juin 2026 · Casablanca", 10.5, GREY, after=10)

p = para(after=4)
runs(p, [("Durée cible : ", {'bold': True}),
         ("≈ 13 min ", {'color': GREEN, 'bold': True}),
         ("(~1800 mots, débit posé d'environ 140 mots/min). Marge prévue pour le créneau "
          "de 15 min. À vous de répartir les 3 rôles entre vous ; les transitions sont "
          "neutres. Chronométrez-vous en répétition et ajustez.", {'color': GREY})])

para("Répartition (à vous de choisir qui prend quel rôle)", 14, DARK, bold=True, before=8, after=4)
tbl = doc.add_table(rows=1, cols=4); tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Light Grid Accent 1'
hdr = tbl.rows[0].cells
for i, h in enumerate(["Rôle", "Parties", "Slides", "Durée"]):
    hdr[i].text = ""; rr = hdr[i].paragraphs[0].add_run(h); rr.font.bold = True
    rr.font.size = Pt(10.5); rr.font.name = FONT
for r in [("Intervenant 1", "P1-P2 — Contexte, objectifs, architecture & conception", "Garde → 06", "≈ 4:40"),
          ("Intervenant 2", "P3-P4 — Application par rôle, intelligence artificielle", "07 → 11", "≈ 4:20"),
          ("Intervenant 3", "P5-P6 — Services, espace patient, conclusion", "12 → 17", "≈ 4:20")]:
    c = tbl.add_row().cells
    for i, v in enumerate(r):
        c[i].text = ""; rr = c[i].paragraphs[0].add_run(v); rr.font.size = Pt(10); rr.font.name = FONT

SPEAKERS = [
("INTERVENANT 1", "Parties 1 & 2 · slides de garde à 06 · ≈ 4 min 40", [
 ("Page de garde", "Accueil & présentation", "0:40",
  "Bonjour Madame et Messieurs les membres du jury. Merci d'être présents aujourd'hui. "
  "Nous sommes Abdou SALOU ABDOU, Kamara MACIRE et Nouridine SAWADOGO, étudiants en "
  "cycle d'ingénieur en informatique et systèmes d'information à l'ISGA. Nous allons "
  "vous présenter CuraMedical : une plateforme intelligente pour gérer un cabinet "
  "médical, avec une aide au diagnostic par intelligence artificielle, la "
  "téléconsultation et l'automatisation du suivi des patients. Ce projet a été réalisé "
  "sous l'encadrement de Dr. Soumia CHOKRI, que nous remercions sincèrement pour son "
  "accompagnement tout au long de ce travail."),
 ("01 / 17", "Plan de la présentation", "0:25",
  "Voici comment va se dérouler notre présentation, en six parties. Le premier "
  "intervenant présentera le contexte, les objectifs, puis l'architecture du système. "
  "Le deuxième présentera l'application selon chaque rôle, ainsi que l'intelligence "
  "artificielle, qui est le cœur du projet. Et le troisième terminera avec les "
  "services avancés, l'espace patient, puis le bilan et les perspectives. Entrons tout "
  "de suite dans le sujet."),
 ("02 / 17", "Contexte & problématique", "0:52",
  "Partons d'un constat simple. Aujourd'hui, dans beaucoup de cabinets, presque tout "
  "se fait encore à la main : dossiers papier, fichiers Excel, agenda manuel et coups "
  "de téléphone. Et cela pose de vrais problèmes. D'abord, le médecin passe trop de "
  "temps sur l'administratif, et donc moins de temps avec ses patients. Ensuite, les "
  "erreurs et les oublis se multiplient : un rendez-vous noté nulle part, un suivi "
  "perdu de vue. À cela s'ajoute un manque important : il n'existe aucun outil pour "
  "aider le médecin dans son diagnostic. Et la communication avec le patient reste "
  "très limitée. C'est tout ce désordre que CuraMedical vient résoudre."),
 ("03 / 17", "Objectifs & périmètre", "0:43",
  "Notre réponse est claire : tout réunir au même endroit, dans une seule application "
  "web. Nous avons six objectifs. Premièrement, centraliser les patients, les "
  "rendez-vous, les consultations et les ordonnances. Deuxièmement, aider le médecin à "
  "poser son diagnostic grâce à l'IA. Troisièmement, sécuriser les données avec un "
  "système de rôles. Quatrièmement, créer automatiquement les documents médicaux. "
  "Cinquièmement, faciliter l'accès aux soins, avec la téléconsultation et un espace "
  "patient. Et enfin, automatiser les rappels et les échanges avec le patient."),
 ("04 / 17", "Architecture technique", "0:58",
  "Pour y arriver, nous avons construit une architecture moderne, découpée en "
  "plusieurs services indépendants, le tout dans des conteneurs Docker — ce qui rend "
  "l'installation simple et fiable. L'interface, ce que voit l'utilisateur, est faite "
  "en React avec Tailwind CSS. Elle communique avec une API construite en Django, "
  "protégée par des jetons de connexion JWT. Point important : nous avons mis "
  "l'intelligence artificielle dans un service à part, en Flask. Comme ça, on peut "
  "l'améliorer sans toucher au reste de l'application. Les données sont enregistrées "
  "dans une base PostgreSQL. Enfin, plusieurs outils sont connectés autour : n8n pour "
  "automatiser, Twilio pour WhatsApp, Jitsi pour la visio, et Groq pour le chatbot."),
 ("05 / 17", "Cas d'utilisation & rôles", "0:38",
  "L'application est organisée autour de quatre rôles, et chacun a ses propres droits. "
  "L'administrateur gère la configuration et les comptes. La secrétaire s'occupe des "
  "patients, des rendez-vous et du planning. Le médecin réalise les consultations, "
  "rédige les ordonnances et utilise l'aide de l'IA. Et le patient a son espace "
  "personnel pour ses rendez-vous et ses documents. L'idée est simple : chacun voit "
  "seulement ce dont il a besoin, ni plus, ni moins."),
 ("06 / 17", "Modèle de données", "0:44",
  "Tout cela repose sur un modèle de données simple et bien organisé, que vous voyez "
  "ici. Les éléments principaux sont : l'Utilisateur, le Patient, le Rendez-vous, la "
  "Consultation et l'Ordonnance. Et les liens entre eux suivent la réalité du terrain : "
  "un patient peut avoir plusieurs rendez-vous et plusieurs consultations, et chaque "
  "consultation donne une ordonnance et un compte-rendu. Le médecin, lui, est un "
  "utilisateur avec un rôle particulier. Ce socle garantit que l'information reste "
  "cohérente du début à la fin. Je laisse maintenant la parole à l'intervenant "
  "suivant, qui va vous montrer l'application en action."),
]),
("INTERVENANT 2", "Parties 3 & 4 · slides 07 à 11 · ≈ 4 min 20", [
 ("07 / 17", "Sécurité & contrôle d'accès", "0:50",
  "Comme on vient de le voir, la sécurité est un point central de CuraMedical, parce "
  "qu'on manipule des données de santé, très sensibles. La connexion se fait avec des "
  "jetons JWT. Mais le plus important, c'est que chaque rôle ne voit que ce qui le "
  "concerne, et cette vérification se fait du côté du serveur, pas seulement à "
  "l'écran. Donc si quelqu'un essaie d'ouvrir une page qui n'est pas pour lui, c'est "
  "bloqué tout de suite — c'est ce que montre la capture à l'écran. Et pour finir, une "
  "corbeille et un historique des actions permettent de revenir en arrière en cas "
  "d'erreur."),
 ("08 / 17", "Gestion patients & rendez-vous", "0:43",
  "Regardons maintenant le quotidien de la secrétaire, qui est en première ligne. "
  "Elle a accès à des dossiers patients complets, avec une recherche instantanée — "
  "fini les classeurs à feuilleter. Elle crée et suit les rendez-vous, qu'elle voit "
  "dans un planning clair, par jour et par médecin. Chaque rendez-vous a un statut : "
  "confirmé, en attente ou annulé. Et si elle supprime quelque chose par erreur, la "
  "corbeille permet de le récupérer. Le but : un accueil rapide, simple, et sans rien "
  "perdre."),
 ("09 / 17", "Consultations (médecin)", "0:44",
  "Côté médecin, tout est pensé pour l'aider à bien décider. Il a sous les yeux une "
  "fiche patient complète : informations, antécédents, allergies et historique des "
  "consultations. Il remplit sa consultation de façon structurée, en indiquant les "
  "symptômes qu'il observe. Et depuis le même écran, il génère directement "
  "l'ordonnance et le compte-rendu. Il a donc une vue d'ensemble, sans passer d'un "
  "outil à l'autre, ce qui lui permet de suivre le patient dans le temps. Et c'est "
  "aussi sur cet écran qu'arrive l'intelligence artificielle, que je vais vous "
  "présenter."),
 ("10 / 17", "Assistance au diagnostic par IA", "1:05",
  "Nous arrivons au cœur du projet : l'aide au diagnostic par intelligence "
  "artificielle. Et c'est très simple à utiliser. Pendant la consultation, le médecin "
  "coche, dans une liste, les symptômes qu'il observe chez son patient. Ces symptômes "
  "partent vers notre service d'IA, qui les transforme en données que le modèle "
  "comprend. Le modèle analyse alors l'ensemble, et renvoie en une fraction de seconde "
  "les maladies les plus probables, classées de la plus probable à la moins probable "
  "— en général les trois premières, avec un score. Et j'insiste sur un point très "
  "important : c'est une aide, pas un diagnostic automatique. Le médecin garde "
  "toujours le dernier mot. L'outil ne décide pas à sa place : il lui donne des "
  "pistes et lui fait gagner du temps."),
 ("11 / 17", "Performances du modèle IA", "0:58",
  "Parlons des résultats, parce qu'une aide au diagnostic n'est utile que si on peut "
  "lui faire confiance. Nous utilisons un modèle appelé forêt aléatoire, entraîné avec "
  "Scikit-learn. Pour l'entraîner, nous avons utilisé quinze mille cas, trois cent "
  "soixante-dix-sept symptômes possibles, et six cent cinquante-cinq maladies. Une "
  "fois entraîné, le modèle donne 77,3 % de bonnes réponses sur des cas qu'il n'avait "
  "jamais vus. Ce chiffre peut paraître moyen, mais il faut le comparer : avec six "
  "cent cinquante-cinq maladies, répondre au hasard donnerait à peine 0,15 %. Notre "
  "modèle fait donc beaucoup mieux, sur une tâche vraiment difficile. Je passe la "
  "parole à l'intervenant suivant."),
]),
("INTERVENANT 3", "Parties 5 & 6 · slides 12 à 17 · ≈ 4 min 20", [
 ("12 / 17", "Ordonnances & documents", "0:47",
  "Poursuivons. Une fois la consultation finie, CuraMedical va jusqu'au bout : il crée "
  "le document tout seul. Le médecin saisit l'ordonnance de façon structurée : pour "
  "chaque médicament, la dose et la durée. À partir de là, l'ordonnance est générée "
  "automatiquement en PDF, avec le cachet et la signature du médecin déjà dessus — le "
  "document est prêt à imprimer ou à envoyer. Les comptes-rendus de consultation "
  "fonctionnent pareil : ils sont aussi exportés en PDF. On passe donc de la saisie à "
  "un document propre et professionnel, en un seul clic, sans tout retaper."),
 ("13 / 17", "Téléconsultation", "0:37",
  "CuraMedical propose aussi la téléconsultation, pour rapprocher le médecin du "
  "patient. La visio est intégrée directement dans la plateforme grâce à Jitsi : rien "
  "à installer. Pour chaque rendez-vous à distance, une salle privée est créée "
  "automatiquement. Et pendant l'appel, le médecin peut prendre ses notes en direct. "
  "C'est très utile pour les patients loin du cabinet, ou qui ont du mal à se "
  "déplacer."),
 ("14 / 17", "Chatbot IA", "0:41",
  "Nous avons aussi ajouté un chatbot, qui permet de poser des questions au cabinet en "
  "langage normal. Concrètement, on tape sa question en français — par exemple "
  "« combien de patients avons-nous ? » ou « quels sont les rendez-vous "
  "d'aujourd'hui ? » — et on a la réponse tout de suite, à partir des vraies données "
  "de l'application. Ce chatbot utilise un grand modèle de langage, via Groq. C'est "
  "donc une deuxième utilisation de l'IA, cette fois pour faciliter la gestion de tous "
  "les jours."),
 ("15 / 17", "Notifications WhatsApp", "0:41",
  "Pour communiquer avec le patient, nous avons automatisé l'envoi de messages par "
  "WhatsApp, l'application que tout le monde utilise. Le patient reçoit des rappels de "
  "rendez-vous automatiques, ce qui évite beaucoup d'oublis, et il peut recevoir son "
  "ordonnance directement sur son téléphone. Tout ça fonctionne grâce à un scénario "
  "automatique fait avec n8n, relié à Twilio. Le résultat est double : moins de "
  "rendez-vous manqués pour le cabinet, et un patient mieux suivi et plus rassuré."),
 ("16 / 17", "Espace patient", "0:39",
  "Le patient, enfin, n'est pas juste un spectateur : il devient acteur de son suivi. "
  "Depuis son espace personnel, il peut s'inscrire en ligne, voir ses prochains "
  "rendez-vous, retrouver ses consultations et ses ordonnances, et modifier ses "
  "informations. Il a donc, à tout moment et où qu'il soit, une vue claire sur son "
  "suivi médical. La boucle est complète : de la secrétaire au médecin, et du médecin "
  "jusqu'au patient."),
 ("17 / 17", "Bilan & perspectives", "0:56",
  "Pour conclure. Nous sommes partis d'un problème concret — un cabinet désorganisé, "
  "sans aide à la décision — et nous avons abouti à une plateforme complète et qui "
  "fonctionne : quatre rôles bien séparés et sécurisés, une intelligence artificielle "
  "réellement opérationnelle, des documents créés automatiquement, et de vraies "
  "connexions avec des outils externes. Ce projet nous a fait vivre toutes les étapes "
  "du métier d'ingénieur, de l'analyse du besoin jusqu'à la mise en service. Et nous "
  "avons des idées pour aller plus loin : mettre la plateforme en production, "
  "continuer à améliorer le modèle d'IA, créer une application mobile pour les "
  "patients, et la relier à d'autres systèmes de santé. Merci beaucoup pour votre "
  "attention. Nous sommes maintenant prêts à répondre à vos questions."),
]),
]

for head, sub, slides in SPEAKERS:
    doc.add_page_break()
    para(head, 16, GREEN, bold=True, after=0)
    para(sub, 10.5, GREY, italic=True, after=10)
    for label, title, timing, speech in slides:
        p = para(before=8, after=3)
        runs(p, [("▸ Slide " + label + " — ", {'bold': True, 'size': 12, 'color': DARK}),
                 (title, {'bold': True, 'size': 12, 'color': DARK}),
                 ("   (≈ " + timing + ")", {'italic': True, 'size': 10, 'color': GREEN})])
        para(speech, 11, DARK, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6, ls=1.18)

doc.save("CuraMedical-Script-Soutenance.docx")
print("OK -> CuraMedical-Script-Soutenance.docx")
