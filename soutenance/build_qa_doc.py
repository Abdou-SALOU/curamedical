# -*- coding: utf-8 -*-
"""Génère un document Word Questions/Réponses sur le fonctionnement de CuraMedical.

Compile les principales questions posées sur le projet, leurs réponses et de
vraies illustrations issues du dépôt (captures d'écran, diagrammes, workflow n8n).
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Chemins ────────────────────────────────────────────────────────────────────
ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
IMG_PRES = os.path.join(ROOT, "soutenance", "presentation", "images")
IMG_SHOT = os.path.join(ROOT, "soutenance", "screenshots", "curamedical-rapport")
OUTPUT = os.path.join(ROOT, "soutenance", "CuraMedical_Questions_Reponses.docx")

# ── Palette ────────────────────────────────────────────────────────────────────
BRAND = RGBColor(0x1F, 0x7A, 0x54)      # vert médical
BRAND_LIGHT = "E8F5EE"                    # fond vert clair
DARK = RGBColor(0x0F, 0x17, 0x2A)
GRAY = RGBColor(0x64, 0x74, 0x8B)
BOX_BG = "F1F5F9"                          # fond gris clair (encadré)
INFO_BG = "FFF7E6"                         # fond ambre (astuce)


def img(folder, name):
    return os.path.join(folder, name)


# ── Helpers de mise en forme ────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def shade_paragraph(p, hex_color):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.space_before = Pt(12)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = BRAND
        p.space_before = Pt(18)
        # filet sous le titre
        pPr = p._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), '1F7A54')
        pbdr.append(bottom)
        pPr.append(pbdr)
    else:
        run.font.size = Pt(12.5)
        run.font.color.rgb = DARK
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    # support **gras** inline
    parts = text.split("**")
    for i, part in enumerate(parts):
        r = p.add_run(part)
        r.font.size = Pt(11)
        if i % 2 == 1:
            r.bold = True
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2)
        parts = it.split("**")
        for i, part in enumerate(parts):
            r = p.add_run(part)
            r.font.size = Pt(11)
            if i % 2 == 1:
                r.bold = True


def add_image(doc, path, width_cm, caption=None):
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY


def add_images_row(doc, paths, width_cm, caption=None):
    """Affiche plusieurs images côte à côte via un tableau invisible."""
    existing = [p for p in paths if os.path.exists(p)]
    if not existing:
        return
    table = doc.add_table(rows=1, cols=len(existing))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, path in zip(table.rows[0].cells, existing):
        cell_p = cell.paragraphs[0]
        cell_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_p.add_run().add_picture(path, width=Cm(width_cm))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(10)
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GRAY


def add_box(doc, title, text, bg=BOX_BG, accent=BRAND):
    """Encadré (1 cellule) avec fond coloré."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    if title:
        r = p.add_run(title + "  ")
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = accent
    parts = text.split("**")
    for i, part in enumerate(parts):
        r = p.add_run(part)
        r.font.size = Pt(10.5)
        if i % 2 == 1:
            r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ── Document ────────────────────────────────────────────────────────────────────
doc = Document()

# Marges
for section in doc.sections:
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

# Police par défaut
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Page de garde ───────────────────────────────────────────────────────────────
logo = img(IMG_SHOT, "Logo_CuraMedical.png")
if os.path.exists(logo):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    p.add_run().add_picture(logo, width=Cm(7))

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(30)
r = t.add_run("CuraMedical")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = BRAND

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Questions / Réponses sur le fonctionnement de l'application")
r.font.size = Pt(15)
r.font.color.rgb = DARK

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Guide de compréhension technique illustré")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = GRAY

doc.add_page_break()

# ── Introduction ────────────────────────────────────────────────────────────────
add_heading(doc, "Introduction", 1)
add_body(doc, "Ce document regroupe les principales questions posées sur le fonctionnement "
              "de l'application CuraMedical, accompagnées de leurs réponses et de véritables "
              "illustrations issues du projet (captures d'écran, diagrammes et workflows). "
              "Il a pour objectif de faciliter la compréhension du système au sein de l'équipe.")
add_image(doc, img(IMG_PRES, "04-architecture-globale.png"), 16,
          "Architecture globale de CuraMedical : frontend React, backend Django, microservice IA, "
          "WhatsApp (Twilio) et automatisations n8n.")

# ── Q1 : ngrok ──────────────────────────────────────────────────────────────────
add_heading(doc, "1. À quoi sert ngrok dans ce projet ?", 1)
add_body(doc, "**ngrok** sert à exposer le backend local sur Internet via une **URL publique en HTTPS**, "
              "le temps du développement.")
add_heading(doc, "Le problème", 2)
add_body(doc, "L'intégration WhatsApp via Twilio fonctionne par *webhook* : quand un patient envoie un "
              "message, Twilio (service externe) doit appeler notre backend. Or, en développement, le "
              "serveur tourne sur localhost — une adresse invisible depuis Internet. Twilio ne peut donc "
              "pas le joindre.")
add_heading(doc, "La solution", 2)
add_bullets(doc, [
    "La commande **ngrok http 8000** crée un tunnel sécurisé vers le serveur local.",
    "Elle génère une URL publique du type **https://xxxx.ngrok.io**.",
    "Cette URL est renseignée dans la variable **PUBLIC_BASE_URL** et déclarée comme webhook dans Twilio.",
])
add_box(doc, "💡 À retenir :",
        "ngrok est un pont temporaire utilisé uniquement en développement. En production, on le remplace "
        "par la vraie URL du serveur (ex. https://api.curamedical.com).", INFO_BG)

# ── Q2 : envoi ordonnances / comptes rendus ─────────────────────────────────────
add_heading(doc, "2. Comment se fait l'envoi des comptes rendus et des ordonnances ?", 1)
add_body(doc, "Les deux suivent le **même schéma** : une action du médecin déclenche une **tâche en "
              "arrière-plan** qui génère un PDF et l'envoie au patient par **plusieurs canaux**.")
add_bullets(doc, [
    "**Génération du PDF** : ordonnance ou compte rendu de consultation.",
    "**Email** : le PDF est joint et envoyé directement au patient.",
    "**WhatsApp** : envoi du PDF via Twilio, à partir de l'URL publique du document.",
    "**n8n** : le PDF (encodé) est transmis au webhook pour d'autres automatisations.",
])
add_body(doc, "Le traitement est **asynchrone** : la requête du médecin rend la main immédiatement, et "
              "l'envoi se fait en tâche de fond. Le médecin n'attend donc jamais.")
add_images_row(doc, [
    img(IMG_SHOT, "38-ordonnance-pdf.png"),
    img(IMG_SHOT, "39-compte-rendu-pdf.png"),
], 7.2, "À gauche : ordonnance générée en PDF. À droite : compte rendu de consultation en PDF.")
add_image(doc, img(IMG_PRES, "15-notifications-whatsapp.jpg"), 9,
          "Exemple de notification reçue par le patient sur WhatsApp.")
add_box(doc, "💡 Point clé :",
        "L'envoi est découplé (asynchrone). C'est pourquoi l'interface reste fluide même si l'email ou "
        "le WhatsApp prend quelques secondes à partir.", INFO_BG)

# ── Q3 : WhatsApp intelligent + diagnostics ─────────────────────────────────────
add_heading(doc, "3. Comment WhatsApp répond-il intelligemment et établit-il les diagnostics ?", 1)
add_body(doc, "Le chatbot WhatsApp s'appuie sur **deux intelligences artificielles distinctes** qui "
              "travaillent ensemble.")
add_heading(doc, "IA n°1 — Comprendre le message (LLaMA via Groq)", 2)
add_body(doc, "Le message du patient est analysé par le modèle **LLaMA 3.3 70B** (servi par Groq), qui "
              "renvoie un résultat structuré : l'intention du patient et les informations extraites "
              "(date, heure, symptômes…). C'est ce qui rend le bot « intelligent » : il comprend une "
              "formulation libre, peu importe la tournure.")
add_heading(doc, "IA n°2 — Le diagnostic (modèle de Machine Learning)", 2)
add_body(doc, "Lorsque le patient décrit des symptômes, ceux-ci sont envoyés au **microservice IA** "
              "(modèle entraîné sur plus de 100 maladies). Celui-ci renvoie les **maladies probables "
              "accompagnées d'un pourcentage de confiance**, que le bot présente en français.")
add_images_row(doc, [
    img(IMG_PRES, "14-chatbot-ia.png"),
    img(IMG_PRES, "10-ia-diagnostic-suggestion.png"),
], 7.6, "Assistant conversationnel et suggestions de diagnostic proposées par l'IA.")
add_images_row(doc, [
    img(IMG_SHOT, "Whatsap_Test_1.jpg"),
    img(IMG_SHOT, "Whatsap_Test_2.jpg"),
    img(IMG_SHOT, "Whatsap_Test_3.jpg"),
], 5.0, "Tests réels du chatbot sur WhatsApp.")
add_box(doc, "⚠️ Sécurité médicale :",
        "LLaMA comprend mais ne pose jamais de diagnostic ; c'est le modèle ML qui calcule les "
        "probabilités. Le résultat est toujours présenté comme indicatif, avec invitation à consulter "
        "un médecin. Les messages vocaux sont d'abord transcrits par Whisper.", INFO_BG)

# ── Q4 : entraînement du modèle ─────────────────────────────────────────────────
add_heading(doc, "4. Quelles sont les étapes d'entraînement du modèle IA ?", 1)
add_body(doc, "Le modèle principal (prédiction de maladies) suit les **étapes classiques du Machine "
              "Learning** :")
add_bullets(doc, [
    "**Charger les données** : un fichier CSV où chaque ligne est un cas (symptômes → maladie).",
    "**Nettoyer** : supprimer les lignes incomplètes et séparer les symptômes (X) de la maladie (y).",
    "**Encoder** : transformer les noms de maladies en nombres (LabelEncoder).",
    "**Diviser** : 90 % des données pour l'entraînement, 10 % pour le test.",
    "**Entraîner** : un modèle Random Forest apprend les liens symptômes → maladies.",
    "**Évaluer** : mesurer la précision (accuracy) sur les données de test.",
    "**Sauvegarder** : enregistrer le modèle en fichiers .pkl réutilisés par l'API.",
])
add_body(doc, "Un **second modèle**, plus léger (TF-IDF + LinearSVC), sert à comprendre les intentions "
              "des phrases du chatbot.")
add_image(doc, img(IMG_SHOT, "33-medecin-consultation-analyse-ia.png"), 16,
          "Le modèle entraîné en action : suggestions de diagnostic lors d'une consultation.")
add_box(doc, "💡 À retenir :",
        "Une fois entraîné et sauvegardé, le modèle est simplement chargé par l'API pour répondre "
        "instantanément, sans ré-entraînement à chaque requête.", INFO_BG)

# ── Q5 : rappel 24h ─────────────────────────────────────────────────────────────
add_heading(doc, "5. Existe-t-il un système de rappel des rendez-vous 24h avant ?", 1)
add_body(doc, "**Oui**, un rappel automatique de la veille existe — implémenté comme **workflow n8n** "
              "(et non dans le backend Django).")
add_bullets(doc, [
    "Un **déclencheur horaire** se lance chaque jour à 8h.",
    "Le workflow récupère **tous les rendez-vous du lendemain**.",
    "Pour chaque RDV, il envoie au patient un **rappel par email ET par WhatsApp** (heure, médecin, motif).",
])
add_box(doc, "⚠️ Important :",
        "Le workflow est livré désactivé par défaut : il faut l'importer dans n8n, configurer les "
        "identifiants (SMTP / Twilio) et l'activer pour qu'il s'exécute.", INFO_BG)

# ── Q6 : ORM ────────────────────────────────────────────────────────────────────
add_heading(doc, "6. Qu'est-ce que l'ORM ?", 1)
add_body(doc, "**ORM = Object-Relational Mapping** (mapping objet-relationnel). C'est la couche qui "
              "**traduit entre les objets Python et les tables de la base de données**, afin de ne "
              "presque jamais écrire de SQL à la main. Dans CuraMedical, il s'agit de l'ORM de Django.")
# Tableau comparatif
table = doc.add_table(rows=4, cols=2)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for c, txt in zip(hdr, ["Sans ORM (SQL brut)", "Avec l'ORM Django"]):
    set_cell_bg(c, BRAND_LIGHT)
    rr = c.paragraphs[0].add_run(txt)
    rr.bold = True
    rr.font.size = Pt(10.5)
rows = [
    ("INSERT INTO patients (...) VALUES (...)", "Patient.objects.create(nom=...)"),
    ("SELECT * FROM patients WHERE id=5", "Patient.objects.get(id=5)"),
    ("SELECT * FROM rendez_vous WHERE patient_id=5", "RendezVous.objects.filter(patient=p)"),
]
for i, (a, b) in enumerate(rows, start=1):
    cells = table.rows[i].cells
    for c, txt in zip(cells, [a, b]):
        rr = c.paragraphs[0].add_run(txt)
        rr.font.size = Pt(10)
        rr.font.name = 'Consolas'
doc.add_paragraph()
add_body(doc, "Les 3 idées clés : une **classe Python = une table**, un **objet = une ligne**, et l'ORM "
              "**génère le SQL automatiquement** (avec protection contre les injections SQL).")

# ── Q7 : temps de réponse ───────────────────────────────────────────────────────
add_heading(doc, "7. Quels sont les temps de réponse ?", 1)
add_body(doc, "Voici des ordres de grandeur estimés d'après l'architecture du système :")
perf = [
    ("Action", "Temps typique"),
    ("Compréhension du chatbot (LLaMA / Groq)", "~ 0,5 – 2 s"),
    ("Diagnostic ML (prédiction des maladies)", "~ 0,1 – 0,5 s"),
    ("Réponse WhatsApp complète (texte)", "~ 1 – 3 s"),
    ("Réponse WhatsApp (message vocal)", "~ 3 – 6 s"),
    ("Envoi ordonnance / compte rendu (côté médecin)", "quasi instantané (< 0,2 s)"),
    ("Livraison réelle email + WhatsApp (arrière-plan)", "~ 2 – 6 s"),
]
ptable = doc.add_table(rows=len(perf), cols=2)
ptable.style = 'Light Grid Accent 1'
ptable.alignment = WD_TABLE_ALIGNMENT.CENTER
for j, (a, b) in enumerate(perf):
    cells = ptable.rows[j].cells
    for c, txt in zip(cells, [a, b]):
        rr = c.paragraphs[0].add_run(txt)
        rr.font.size = Pt(10.5)
        if j == 0:
            rr.bold = True
            set_cell_bg(c, BRAND_LIGHT)
doc.add_paragraph()
add_box(doc, "💡 Point clé :",
        "Pour les ordonnances et comptes rendus, l'envoi est asynchrone : le médecin n'attend jamais "
        "que l'email ou le WhatsApp parte. Les temps dépendent surtout des services externes (Groq, "
        "Twilio, SMTP) et du réseau.", INFO_BG)

# ── Q8 : workflows n8n ──────────────────────────────────────────────────────────
add_heading(doc, "8. À quoi servent les workflows n8n ?", 1)
add_body(doc, "n8n automatise **toutes les communications patients**. Le workflow maître regroupe "
              "**trois automatisations** :")
add_bullets(doc, [
    "**Rappel 24h** : chaque jour à 8h, rappel des RDV du lendemain (email + WhatsApp).",
    "**RDV terminé** : message de suivi envoyé après la clôture d'une consultation.",
    "**Notifications** : un aiguilleur (Switch) envoie le compte-rendu, l'ordonnance ou un message "
    "de bienvenue selon l'événement reçu du backend.",
])
add_image(doc, img(IMG_PRES, "17-workflow-n8n-rogne.png"), 13,
          "Le workflow n8n complet : rappel 24h (haut), RDV terminé (milieu) et notifications (bas).")
add_body(doc, "La logique est commune aux trois : un **déclencheur** (une horloge ou un webhook), un "
              "**aiguillage** des données, la **reconstruction du PDF**, puis l'envoi par **double "
              "canal : email et WhatsApp**.")
add_box(doc, "💡 En une phrase :",
        "Le backend envoie un événement à n8n, qui lit son type et l'achemine vers le bon canal de "
        "communication — chaque message partant à la fois par email et par WhatsApp.", INFO_BG)

# ── Q9 : authentification & sécurité ────────────────────────────────────────────
add_heading(doc, "9. Comment gérez-vous l'authentification et la sécurité ?", 1)
add_body(doc, "L'authentification repose sur des **jetons JWT** (JSON Web Tokens). À la connexion, "
              "l'utilisateur reçoit un jeton d'accès et un jeton de rafraîchissement ; chaque requête "
              "est ensuite authentifiée par ce jeton.")
add_bullets(doc, [
    "Connexion possible par **nom d'utilisateur OU adresse email**.",
    "**Limitation du nombre de tentatives** pour protéger contre les attaques par force brute.",
    "**Quatre rôles** : administrateur, médecin, secrétaire et patient.",
    "**Contrôle d'accès par rôle (RBAC)** : chaque utilisateur ne voit que les données qui le concernent.",
])
add_body(doc, "Par exemple, un patient n'accède qu'à son propre dossier, un médecin à ses patients, et "
              "l'administrateur n'a pas accès aux données médicales sensibles.")
add_images_row(doc, [
    img(IMG_PRES, "07-controle-acces-role.png"),
    img(IMG_SHOT, "25-secretaire-acces-refuse.png"),
], 7.6, "Contrôle d'accès par rôle : chaque profil dispose de permissions précises ; "
        "tout accès non autorisé est refusé.")

# ── Q10 : architecture microservices ────────────────────────────────────────────
add_heading(doc, "10. Pourquoi une architecture en plusieurs services ?", 1)
add_body(doc, "L'application est découpée en **services indépendants**, chacun avec une responsabilité "
              "claire :")
add_bullets(doc, [
    "**Frontend React** : l'interface utilisateur.",
    "**Backend Django** : la logique métier, la base de données et la sécurité.",
    "**Microservice IA (Python/Flask)** : l'intelligence artificielle (diagnostic et chatbot).",
    "**n8n** : les automatisations (emails, WhatsApp, rappels).",
])
add_body(doc, "Cette séparation permet de faire **évoluer ou redémarrer l'IA sans toucher au backend**, "
              "d'utiliser la technologie la plus adaptée à chaque besoin, et de mieux **répartir la "
              "charge**.")
add_image(doc, img(IMG_SHOT, "diagramme/architecture_globale.png"), 16,
          "Architecture en services de CuraMedical et communication entre les composants.")

# ── Q11 : téléconsultation ──────────────────────────────────────────────────────
add_heading(doc, "11. Comment fonctionne la téléconsultation vidéo ?", 1)
add_body(doc, "La visioconférence s'appuie sur **Jitsi** (serveur public meet.jit.si) via le SDK React "
              "officiel — **sans aucune clé d'API**.")
add_bullets(doc, [
    "À chaque rendez-vous en ligne, le backend génère un **nom de salle unique** (ex. CuraMedical-RDV-xxxxxxxx).",
    "Le **médecin et le patient** rejoignent cette même salle à l'heure du rendez-vous.",
    "Aucune installation n'est nécessaire : tout se passe directement dans le navigateur.",
])
add_images_row(doc, [
    img(IMG_PRES, "13-teleconsultation.png"),
    img(IMG_SHOT, "37-medecin-teleconsultation-salle-notes.png"),
], 7.6, "Salle de téléconsultation : vidéo entre le médecin et le patient, avec prise de notes.")

# ── Q12 : confidentialité des données ───────────────────────────────────────────
add_heading(doc, "12. Comment protégez-vous les données médicales ?", 1)
add_body(doc, "La protection des données repose sur plusieurs mécanismes complémentaires :")
add_bullets(doc, [
    "**Accès restreint par rôle** : les données médicales ne sont visibles que par le personnel concerné.",
    "**Traçabilité (journal d'audit)** : les modifications sur les dossiers sont enregistrées.",
    "**Suppression douce (corbeille)** : rien n'est effacé définitivement d'un seul clic ; "
    "les éléments peuvent être restaurés.",
    "**Communication chiffrée (HTTPS)** et authentification par jeton pour chaque requête.",
])
add_image(doc, img(IMG_SHOT, "14-medecin-rdv-corbeille.png"), 16,
          "Corbeille : les éléments supprimés sont conservés et peuvent être restaurés, "
          "évitant toute perte de données.")

# ── Q13 : fiabilité de l'IA ─────────────────────────────────────────────────────
add_heading(doc, "13. Quelle est la fiabilité de l'IA et qui est responsable du diagnostic ?", 1)
add_body(doc, "L'IA est une **aide à la décision**, et non un médecin. Elle propose des hypothèses "
              "accompagnées d'un **pourcentage de confiance**, mais ne pose jamais de diagnostic "
              "définitif.")
add_bullets(doc, [
    "Le **médecin garde toujours le contrôle** : il valide, modifie ou écarte les suggestions.",
    "Le résultat est explicitement présenté comme **indicatif**.",
    "Double garde-fou : le modèle de langage comprend mais ne diagnostique pas ; le modèle de "
    "Machine Learning calcule des probabilités ; **c'est le médecin qui tranche**.",
])
add_image(doc, img(IMG_SHOT, "16-medecin-nouvelle-consultation-ia.png"), 16,
          "Le médecin reste décisionnaire : l'IA propose, le praticien dispose.")

# ── Q14 : déploiement ───────────────────────────────────────────────────────────
add_heading(doc, "14. Comment l'application est-elle déployée ?", 1)
add_body(doc, "Toute l'application est **conteneurisée avec Docker** et orchestrée par Docker Compose. "
              "L'écosystème complet se lance avec une seule commande, de façon **reproductible** sur "
              "n'importe quelle machine.")
add_body(doc, "Les services conteneurisés sont :")
add_bullets(doc, [
    "**PostgreSQL** — la base de données",
    "**Redis** — la file d'attente des tâches",
    "**Backend Django** + **worker Celery** — l'API et les tâches en arrière-plan",
    "**Frontend React** — l'interface",
    "**Microservice IA** — l'intelligence artificielle",
    "**n8n** — les automatisations",
])
add_image(doc, img(ROOT, "soutenance/Diagramme_uml/diagramme de déploiement.png"), 15,
          "Diagramme de déploiement : organisation des services de l'application.")

# ── Q15 : robustesse ────────────────────────────────────────────────────────────
add_heading(doc, "15. Que se passe-t-il si un service externe tombe en panne ?", 1)
add_body(doc, "Le système est conçu pour **rester fonctionnel même en cas de défaillance** d'un "
              "composant :")
add_bullets(doc, [
    "Si la file de tâches (Redis/Celery) est indisponible, les notifications **basculent "
    "automatiquement sur un autre mécanisme** : aucune notification n'est perdue et la requête "
    "n'est jamais bloquée.",
    "Si le service de langage (Groq) est indisponible, le chatbot **bascule sur un raisonnement "
    "de secours**.",
    "Grâce à la suppression douce, **aucune donnée n'est perdue définitivement**.",
])
add_box(doc, "💡 À retenir :",
        "La robustesse repose sur des mécanismes de repli (fallback) : l'application continue de "
        "fonctionner, en mode dégradé si nécessaire, plutôt que de s'arrêter.", INFO_BG)

# ── Q16 : provenance du jeu de données ──────────────────────────────────────────
add_heading(doc, "16. D'où vient votre jeu de données et est-il représentatif ?", 1)
add_body(doc, "Le modèle de diagnostic est entraîné sur un **jeu de données public issu de Kaggle** : "
              "« **Diseases and Symptoms Dataset** » (auteur *dhivyeshrk*), dont le fichier d'origine est "
              "**Final_Augmented_dataset_Diseases_and_Symptoms.csv**.")
add_bullets(doc, [
    "Format : une ligne = un cas ; **377 colonnes de symptômes** (présent / absent, 0 ou 1) → **une maladie**.",
    "Dataset complet : **246 945 lignes** et **773 maladies**.",
    "Nous en avons extrait **Fast_Dataset.csv**, un **sous-ensemble de 15 000 lignes (655 maladies)** "
    "pour un entraînement rapide et reproductible.",
])
add_box(doc, "🔗 Source :",
        "Kaggle — « Diseases and Symptoms Dataset » (dhivyeshrk) : "
        "kaggle.com/datasets/dhivyeshrk/diseases-and-symptoms-dataset", INFO_BG)
add_heading(doc, "Est-il représentatif ?", 2)
add_body(doc, "**Avantage** : une couverture très large (plusieurs centaines de maladies). "
              "**Limite assumée** : ce sont des données publiques, agrégées et « augmentées » "
              "(combinaisons symptômes-maladies), et **non de vrais dossiers cliniques** d'un cabinet ; "
              "les symptômes sont en anglais et certaines classes sont déséquilibrées.")
add_box(doc, "⚠️ Honnêteté scientifique :",
        "En production réelle, il faudrait **réentraîner le modèle sur des données du cabinet, validées "
        "par des médecins**. C'est précisément pour cela que l'IA reste une **aide au diagnostic**, "
        "jamais un décideur.", INFO_BG)

# ── Q17 : choix du Random Forest ────────────────────────────────────────────────
add_heading(doc, "17. Pourquoi un Random Forest plutôt que du deep learning ou un LLM ?", 1)
add_body(doc, "Le choix du modèle est dicté par la **nature des données** : des symptômes binaires "
              "(présent / absent), donc des **données tabulaires**.")
add_bullets(doc, [
    "Sur ce type de données, les **forêts aléatoires** sont parmi les plus performantes, **rapides à "
    "entraîner** et **peu gourmandes** (pas besoin de GPU).",
    "Elles sont **explicables** : on peut connaître l'importance de chaque symptôme dans la décision.",
    "Un **réseau de neurones profond** exigerait beaucoup plus de données et serait une **boîte noire**, "
    "peu adaptée à un contexte médical où il faut pouvoir justifier.",
    "Un **LLM** (type LLaMA) n'est pas fait pour ce calcul de probabilités : chez nous, il sert au "
    "**chatbot** (compréhension du langage), pas au diagnostic chiffré.",
])
add_box(doc, "💡 En une phrase :",
        "Le bon outil pour le bon problème — Random Forest pour le diagnostic tabulaire, LLM pour le "
        "langage naturel.", INFO_BG)

# ── Q18 : explicabilité ─────────────────────────────────────────────────────────
add_heading(doc, "18. Le médecin comprend-il pourquoi l'IA propose telle maladie ?", 1)
add_body(doc, "Oui, dans une certaine mesure — c'est justement un avantage du Random Forest.")
add_bullets(doc, [
    "L'IA renvoie les **3 maladies les plus probables**, chacune avec un **score de confiance** (%).",
    "Ces suggestions découlent directement des **symptômes cochés** par le médecin pendant la consultation.",
    "On peut exposer l'**importance des symptômes** : le modèle sait quelles variables pèsent le plus.",
])
add_body(doc, "Le médecin **recoupe toujours** ces pistes avec l'examen clinique : l'outil oriente, "
              "il ne remplace pas le raisonnement.")

# ── Q19 : RGPD / loi 09-08 (CNDP) ───────────────────────────────────────────────
add_heading(doc, "19. Êtes-vous conformes au RGPD / à la loi 09-08 (CNDP) sur les données de santé ?", 1)
add_body(doc, "Les données de santé sont **sensibles** ; le projet intègre plusieurs garde-fous techniques :")
add_bullets(doc, [
    "**Cloisonnement par rôle côté serveur** : un patient ne voit que son propre dossier.",
    "**Journal d'audit** des accès et modifications, **mots de passe hachés**, **jetons JWT**, **HTTPS**.",
    "**Suppression douce** (corbeille) : pas d'effacement définitif accidentel.",
])
add_box(doc, "⚠️ Honnêteté :",
        "Le projet est un **prototype académique**. Pour une mise en production réelle au Maroc, il "
        "faudrait : **déclaration à la CNDP (loi 09-08)**, **consentement explicite du patient**, "
        "**chiffrement des données au repos** et **hébergement agréé** pour les données de santé "
        "(RGPD pour un déploiement en Europe).", INFO_BG)

# ── Q20 : sécurité du canal WhatsApp ────────────────────────────────────────────
add_heading(doc, "20. N'importe qui peut-il écrire au WhatsApp du cabinet et obtenir des informations ?", 1)
add_body(doc, "C'est un **point de vigilance** que nous identifions clairement.")
add_bullets(doc, [
    "Le chatbot répond au **numéro émetteur** ; les actions sensibles (rendez-vous, données patient) "
    "doivent être protégées par une **vérification d'identité**.",
    "Amélioration prévue : **lier le numéro WhatsApp à un compte patient vérifié** avant de divulguer "
    "la moindre information personnelle.",
])
add_body(doc, "Aujourd'hui, le bot reste surtout **informatif et transactionnel** : aucune donnée médicale "
              "détaillée n'est renvoyée sans rattachement à un compte.")

# ── Q21 : tests ─────────────────────────────────────────────────────────────────
add_heading(doc, "21. Comment avez-vous testé l'application ?", 1)
add_bullets(doc, [
    "**Tests fonctionnels manuels** pour chacun des 4 rôles (administrateur, secrétaire, médecin, patient).",
    "Vérification des **endpoints de l'API** et des **accès interdits** (un rôle ne peut pas atteindre "
    "les données d'un autre).",
    "Tests **bout en bout** des parcours clés : prise de rendez-vous, consultation, génération et envoi "
    "des documents, chatbot WhatsApp.",
])
add_box(doc, "🔭 Perspective :",
        "Axe d'amélioration assumé : ajouter des **tests automatisés** (pytest côté Django, Jest côté "
        "React) et une **intégration continue** pour sécuriser les évolutions.", INFO_BG)

# ── Q22 : organisation de l'équipe ──────────────────────────────────────────────
add_heading(doc, "22. Comment avez-vous réparti le travail à trois et géré le code ?", 1)
add_body(doc, "Le travail a été découpé en **briques cohérentes**, chacun pilotant un domaine tout en "
              "collaborant sur l'ensemble :")
add_bullets(doc, [
    "**Backend & IA** : Django, microservice de diagnostic, entraînement du modèle.",
    "**Frontend** : React, interfaces des 4 rôles et espace patient.",
    "**Automatisations & intégrations** : n8n, WhatsApp/Twilio, téléconsultation, déploiement Docker.",
])
add_body(doc, "Le code est versionné avec **Git/GitHub** (branches, commits, historique), avec des "
              "**points réguliers** pour intégrer les différentes parties.")

# ── Q23 : montée en charge ──────────────────────────────────────────────────────
add_heading(doc, "23. L'application peut-elle monter en charge (clinique, plusieurs médecins) ?", 1)
add_body(doc, "L'architecture a été pensée pour **évoluer** :")
add_bullets(doc, [
    "Les services sont **séparés et conteneurisés (Docker)** → on peut **redémarrer ou multiplier** "
    "le backend et le service IA indépendamment.",
    "**PostgreSQL** gère sans difficulté un volume important de dossiers et de rendez-vous.",
    "Les tâches lourdes (emails, PDF, WhatsApp) sont **asynchrones** (Celery/Redis) → elles ne "
    "ralentissent pas l'interface.",
])
add_box(doc, "💡 Limite actuelle :",
        "Le projet est dimensionné pour **un cabinet** (mono-instance). Passer à une clinique "
        "multi-cabinets demanderait surtout de la configuration et de la mise à l'échelle, pas une "
        "réécriture.", INFO_BG)

# ── Q24 : différenciation ───────────────────────────────────────────────────────
add_heading(doc, "24. Qu'est-ce qui vous distingue des solutions existantes (Doctolib, etc.) ?", 1)
add_body(doc, "Les solutions connues (type Doctolib) sont surtout des **plateformes de prise de "
              "rendez-vous**. CuraMedical va plus loin :")
add_bullets(doc, [
    "**Tout-en-un** : gestion du cabinet **+ aide au diagnostic par IA + téléconsultation + "
    "automatisation WhatsApp**, dans un seul outil.",
    "**Auto-hébergeable** et open source : le cabinet **garde la maîtrise de ses données**.",
    "**Adapté au contexte local** : petit cabinet, communication par WhatsApp que tout le monde utilise.",
])
add_body(doc, "Notre différenciateur principal reste l'**aide au diagnostic intégrée à la consultation**, "
              "absente des agendas médicaux classiques.")

# ── Q25 : évolution du modèle ───────────────────────────────────────────────────
add_heading(doc, "25. Comment le modèle d'IA évolue-t-il dans le temps ?", 1)
add_body(doc, "Le modèle d'IA est **isolé dans un microservice**, ce qui simplifie ses mises à jour :")
add_bullets(doc, [
    "Le modèle est entraîné par un **script séparé** (train.py) qui produit des fichiers **.pkl**.",
    "Pour le mettre à jour, on **réentraîne et on remplace ces fichiers** — **sans toucher** au "
    "backend ni au frontend.",
    "En perspective : une **boucle de retour** où les diagnostics validés par les médecins "
    "**enrichissent les données** d'entraînement.",
])
add_box(doc, "💡 À retenir :",
        "C'est tout l'intérêt d'avoir séparé l'IA : on l'améliore en continu sans interrompre le reste "
        "de l'application.", INFO_BG)

# ── Sauvegarde ──────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Document généré : {OUTPUT}")
