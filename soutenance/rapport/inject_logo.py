# -*- coding: utf-8 -*-
"""Injecte l'icone CuraMedical (page de garde + pied de page) dans le .docx existant
via python-docx (sans Word), en preservant le contenu/retouches manuelles.
Lit toujours la sauvegarde pristine (SRC) et ecrit le livrable (OUT) -> reexecution propre.
Usage: python inject_logo.py [cover_cm]
"""
import sys, os, shutil
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = r"c:\2CI-ISI\S2\Projet en Systèmes Informatiques\MedPredict\CuraMedical\soutenance\screenshots\curamedical-rapport"
SRC  = os.path.join(BASE, "Rapport_CuraMedical.EDITED_20260614_211436.docx")   # version editee pristine
OUT  = os.path.join(BASE, "Rapport_CuraMedical.docx")
LOGO = os.path.join(BASE, "Logo_CuraMedical.png")

cover_cm  = float(sys.argv[1]) if len(sys.argv) > 1 else 2.1
footer_cm = 0.42

assert os.path.exists(SRC),  f"SRC introuvable: {SRC}"
assert os.path.exists(LOGO), f"LOGO introuvable: {LOGO}"

doc = Document(SRC)


def has_drawing(p):
    return bool(p._p.findall('.//' + qn('w:drawing')))


# ── 1) Page de garde : icone au-dessus du titre (apres « RAPPORT DE PROJET ») ──
paras = doc.paragraphs
ti = None
for i, p in enumerate(paras):
    if p.text.strip().upper() == "RAPPORT DE PROJET":
        ti = i + 1
        break

cover_status = "ANCHOR_NOT_FOUND"
if ti is not None and ti < len(paras):
    title = paras[ti]
    if ti >= 1 and has_drawing(paras[ti - 1]):
        cover_status = "ALREADY_PRESENT"
    else:
        newp = title.insert_paragraph_before()
        newp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        newp.paragraph_format.space_before = None
        newp.add_run().add_picture(LOGO, height=Cm(cover_cm))
        cover_status = f"ADDED (title='{title.text.strip()}', {cover_cm}cm)"


# ── 2) Pied de page : petite icone a droite, juste avant « Page » ──
def inject_footer(footer):
    for para in footer.paragraphs:
        target = next((r for r in para.runs if 'Page' in (r.text or '')), None)
        if target is None:
            continue
        if has_drawing(para):
            return "ALREADY_PRESENT"
        pic = para.add_run()
        pic.add_picture(LOGO, height=Cm(footer_cm))
        sp = para.add_run("  ")
        target._r.addprevious(pic._r)
        target._r.addprevious(sp._r)
        return "ADDED"
    return "NO_PAGE_TEXT"


foot_status = []
for si, sec in enumerate(doc.sections):
    f = sec.footer
    if f.is_linked_to_previous:
        foot_status.append(f"sec{si}:LINKED")
        continue
    foot_status.append(f"sec{si}:" + inject_footer(f))

doc.save(OUT)
print("COVER :", cover_status)
print("FOOTER:", ", ".join(foot_status))
print("SAVED :", OUT, "(", round(os.path.getsize(OUT) / 1024, 1), "KB )")
