# -*- coding: utf-8 -*-
"""Orchestration : page de garde, sections, en-tête/pied, assemblage du rapport."""
import os
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from engine import (
    Report, LOGO, LOGO_APP, PRIMARY, PRIMARY_DARK, ACCENT, INK, GREY, WHITE,
    HEX_PRIMARY, HEX_ACCENT, HEAD_FONT, BODY_FONT, MONO_FONT,
    _field, enable_update_fields, set_page_number_start, set_borders,
    para_shading, _spacing,
)
from content_front import front_matter, introduction
from content_chapters import chapter1, chapter2, chapter3, chapter4, chapter5, chapter6
from content_end import conclusion, bibliographie, annexes, resume, table_des_matieres

OUT = r"c:\2CI-ISI\S2\Projet en Systèmes Informatiques\MedPredict\CuraMedical\soutenance\screenshots\curamedical-rapport\Rapport_CuraMedical.docx"


def build_cover(r):
    doc = r.doc
    sec = doc.sections[0]
    r._format_section(sec)

    # Bandeau supérieur fin (filet vert)
    top = doc.add_paragraph()
    top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    top.paragraph_format.space_before = Pt(6)
    top.paragraph_format.space_after = Pt(4)
    set_borders(top, {'bottom': ('single', 12, 6, HEX_PRIMARY)})

    # Logo ISGA
    if os.path.exists(LOGO):
        plogo = doc.add_paragraph(); plogo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        plogo.paragraph_format.space_before = Pt(6)
        plogo.paragraph_format.space_after = Pt(2)
        plogo.add_run().add_picture(LOGO, width=Cm(4.6))

    # Établissement
    def line(text, size, color, bold=False, italic=False, before=0, after=4, font=HEAD_FONT, spacing=None, caps=False):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        rr = p.add_run(text.upper() if caps else text)
        rr.font.name = font; rr.font.size = Pt(size); rr.font.bold = bold; rr.font.italic = italic
        rr.font.color.rgb = color
        if spacing: _spacing(rr, spacing)
        return p

    line("Institut Supérieur d'Ingénierie et des Affaires — ISGA", 13, PRIMARY_DARK, bold=True, before=2, after=1)
    line("Cycle d'Ingénierie · Informatique & Systèmes d'Information (2CI-ISI)", 11, INK, after=1)
    line("Projet en Systèmes Informatiques — Groupe 2 · Année académique 2025–2026", 10.5, GREY, italic=True, after=2)

    # Filet
    sep = doc.add_paragraph(); sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.space_before = Pt(6); sep.paragraph_format.space_after = Pt(14)
    set_borders(sep, {'bottom': ('single', 6, 4, "BFE3CE")})

    # Sur-titre
    line("RAPPORT DE PROJET", 12, ACCENT, bold=True, spacing=80, after=8, caps=True)

    # Logo (icône de marque) de l'application CuraMedical
    if os.path.exists(LOGO_APP):
        papp = doc.add_paragraph(); papp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        papp.paragraph_format.space_before = Pt(2)
        papp.paragraph_format.space_after = Pt(4)
        papp.add_run().add_picture(LOGO_APP, height=Cm(2.3))

    # Titre principal
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    rr = p.add_run("CuraMedical")
    rr.font.name = HEAD_FONT; rr.font.size = Pt(40); rr.font.bold = True; rr.font.color.rgb = PRIMARY_DARK

    # Sous-titre produit
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    rr = p.add_run("Plateforme intelligente de gestion de cabinet médical")
    rr.font.name = HEAD_FONT; rr.font.size = Pt(16); rr.font.bold = True; rr.font.color.rgb = PRIMARY
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    rr = p.add_run("Aide au diagnostic par intelligence artificielle, téléconsultation et automatisation du parcours de soins")
    rr.font.name = BODY_FONT; rr.font.size = Pt(11.5); rr.font.italic = True; rr.font.color.rgb = GREY

    # Filet décoratif
    sep = doc.add_paragraph(); sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.space_before = Pt(2); sep.paragraph_format.space_after = Pt(18)
    set_borders(sep, {'bottom': ('single', 12, 4, HEX_PRIMARY)})

    # Bloc auteurs / encadrante (tableau 2 colonnes sans bordure visible)
    from docx.enum.table import WD_TABLE_ALIGNMENT
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c1, c2 = tbl.rows[0].cells
    c1.width = Cm(9.0); c2.width = Cm(7.0)

    def block(cell, title, names):
        cell.paragraphs[0].text = ""
        pt = cell.paragraphs[0]
        pt.paragraph_format.space_after = Pt(4)
        rt = pt.add_run(title)
        rt.font.name = HEAD_FONT; rt.font.size = Pt(11); rt.font.bold = True; rt.font.color.rgb = PRIMARY
        _spacing(rt, 30)
        for n in names:
            pn = cell.add_paragraph(); pn.paragraph_format.space_after = Pt(2)
            rn = pn.add_run(n)
            rn.font.name = BODY_FONT; rn.font.size = Pt(12); rn.font.color.rgb = INK

    block(c1, "RÉALISÉ PAR", ["Abdou SALOU ABDOU", "Kamara MACIRE", "Nouridine SAWADOGO"])
    block(c2, "ENCADRÉ PAR", ["Dr. Soumia CHOKRI"])

    # Pied de page de garde (date/lieu)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    foot = doc.add_paragraph(); foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.paragraph_format.space_before = Pt(20)
    set_borders(foot, {'top': ('single', 6, 6, "BFE3CE")})
    rf = foot.add_run("Soutenance — Juin 2026 · Casablanca, Maroc")
    rf.font.name = HEAD_FONT; rf.font.size = Pt(10.5); rf.font.color.rgb = GREY; rf.font.italic = True


def setup_body_section(r):
    """Crée la 2e section (corps), numérotée à partir de 1, avec en-tête/pied."""
    doc = r.doc
    new = doc.add_section(WD_SECTION.NEW_PAGE)
    r._format_section(new)
    new.different_first_page_header_footer = False
    set_page_number_start(new, 1)

    # délier de la section précédente
    new.header.is_linked_to_previous = False
    new.footer.is_linked_to_previous = False

    # ── En-tête : nom du projet à gauche, tagline à droite, filet vert ──
    hp = new.header.paragraphs[0]
    hp.text = ""
    tabs = hp.paragraph_format.tab_stops
    tabs.add_tab_stop(Cm(16.2), WD_TAB_ALIGNMENT.RIGHT)
    rl = hp.add_run("CuraMedical")
    rl.font.name = HEAD_FONT; rl.font.size = Pt(9); rl.font.bold = True; rl.font.color.rgb = PRIMARY
    rr = hp.add_run("\tPlateforme intelligente de gestion de cabinet médical")
    rr.font.name = HEAD_FONT; rr.font.size = Pt(8.5); rr.font.italic = True; rr.font.color.rgb = GREY
    set_borders(hp, {'bottom': ('single', 6, 6, "BFE3CE")})

    # ── Pied : établissement à gauche, "Page X" à droite ──
    fp = new.footer.paragraphs[0]
    fp.text = ""
    tabs = fp.paragraph_format.tab_stops
    tabs.add_tab_stop(Cm(16.2), WD_TAB_ALIGNMENT.RIGHT)
    set_borders(fp, {'top': ('single', 6, 6, "BFE3CE")})
    rl = fp.add_run("ISGA · 2CI-ISI — Groupe 2")
    rl.font.name = HEAD_FONT; rl.font.size = Pt(8.5); rl.font.color.rgb = GREY
    rmid = fp.add_run("\t")
    # Icône de marque CuraMedical, alignée à droite, juste avant le numéro de page
    if os.path.exists(LOGO_APP):
        fp.add_run().add_picture(LOGO_APP, height=Cm(0.42))
        fp.add_run("  ").font.size = Pt(9)
    rp = fp.add_run("Page ")
    rp.font.name = HEAD_FONT; rp.font.size = Pt(9); rp.font.color.rgb = PRIMARY_DARK; rp.font.bold = True
    _field(fp, "PAGE", cached="1", font=HEAD_FONT, size=9, color=PRIMARY_DARK, bold=True)


def main():
    r = Report()
    build_cover(r)
    setup_body_section(r)

    front_matter(r)
    introduction(r)
    chapter1(r)
    chapter2(r)
    chapter3(r)
    chapter4(r)
    chapter5(r)
    chapter6(r)
    conclusion(r)
    bibliographie(r)
    annexes(r)
    table_des_matieres(r)   # TDM détaillée en fin de document (style mémoire)
    resume(r)               # Résumé en toute fin

    # NB : pas de updateFields à l'ouverture (évite le pop-up Word).
    # La TDM / listes / numéros de page sont figés ensuite via l'automatisation Word.
    r.save(OUT)
    print("OK ->", OUT)


if __name__ == '__main__':
    main()
