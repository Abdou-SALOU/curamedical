# -*- coding: utf-8 -*-
import os
from collections import Counter
from docx import Document

p = r'soutenance/screenshots/curamedical-rapport/Rapport_CuraMedical.docx'
d = Document(p)
body = d.element.body.xml
print('Taille:', round(os.path.getsize(p) / 1024 / 1024, 1), 'Mo')
print('Placeholder TDM restant ?', ("affichera ici" in body))
print('Hyperliens PAGEREF (entrees TDM/listes):', body.count('PAGEREF'))
print('Images inline:', len(d.inline_shapes))
c = Counter(par.style.name for par in d.paragraphs)
print('H1/H2/H3:', c.get('Heading 1'), c.get('Heading 2'), c.get('Heading 3'))
print('Tables:', len(d.tables))
