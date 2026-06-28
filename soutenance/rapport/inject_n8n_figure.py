# -*- coding: utf-8 -*-
"""Injecte la figure du workflow n8n (capture de l'editeur) dans la section
V.1 « Workflow n8n pour les rappels automatiques » du rapport, via python-docx,
en preservant toutes les retouches manuelles (injection ciblee, pas de regen).

- idempotent : ne reinsere pas si une image est deja presente dans la sous-section ;
- numero de figure calcule en comptant les champs « SEQ Figure » qui precedent ;
- legende au format maison (champ SEQ Figure -> reprise dans la Liste des figures) ;
- sauvegarde horodatee avant ecriture.

Usage : python inject_n8n_figure.py
"""
import os
import shutil
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.dirname(ROOT)
BASE = os.path.join(REPO, "screenshots", "curamedical-rapport")
IMG = os.path.join(REPO, "presentation", "images", "17-workflow-n8n.png")

TARGETS = [
    os.path.join(BASE, "Rapport_CuraMedical.docx"),                      # livrable
    os.path.join(BASE, "Rapport_CuraMedical.EDITED_20260614_211436.docx"),  # pristine (pipeline)
]

HEAD_FONT = "Calibri"
PRIMARY = RGBColor(0x14, 0x7A, 0x4E)
GREY = RGBColor(0x55, 0x63, 0x6B)

ANCHOR_KEY = "workflow n8n pour les rappels"
CAPTION = ("Workflow n8n d'automatisation : rappels J-1, comptes rendus et messages "
           "de bienvenue diffuses par e-mail et WhatsApp (capture de l'editeur n8n)")
WIDTH_CM = 15.5


def _field(paragraph, instr, cached="", font=None, size=None, color=None, bold=None):
    """Insere un champ Word complexe (ici SEQ Figure)."""
    run = paragraph.add_run()
    b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin'); run._r.append(b)
    run2 = paragraph.add_run()
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = instr
    run2._r.append(it)
    run3 = paragraph.add_run()
    s = OxmlElement('w:fldChar'); s.set(qn('w:fldCharType'), 'separate'); run3._r.append(s)
    run4 = paragraph.add_run(cached)
    run5 = paragraph.add_run()
    e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end'); run5._r.append(e)
    if font:  run4.font.name = font
    if size:  run4.font.size = Pt(size)
    if color: run4.font.color.rgb = color
    if bold is not None: run4.font.bold = bold
    return run4


def keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(OxmlElement('w:keepNext'))


def has_drawing(p):
    return bool(p._p.findall('.//' + qn('w:drawing')))


def count_seq_figure(paragraphs, upto):
    n = 0
    for p in paragraphs[:upto]:
        for it in p._p.findall('.//' + qn('w:instrText')):
            if it.text and 'SEQ Figure' in it.text:
                n += 1
    return n


def inject(path):
    if not os.path.exists(path):
        return f"SKIP (introuvable) {os.path.basename(path)}"
    doc = Document(path)
    paras = doc.paragraphs

    # 1) ancre : titre Heading de la sous-section V.1
    head = None
    for i, p in enumerate(paras):
        if ANCHOR_KEY in p.text.strip().lower() and p.style.name.lower().startswith("heading"):
            head = i
            break
    if head is None:
        return f"ANCHOR_NOT_FOUND {os.path.basename(path)}"

    # 2) fin de la sous-section = prochain titre Heading
    anchor = None
    for j in range(head + 1, len(paras)):
        if paras[j].style.name.lower().startswith("heading"):
            anchor = j
            break
    if anchor is None:
        anchor = len(paras)

    # 3) idempotence : image deja presente dans la sous-section ?
    for p in paras[head + 1:anchor]:
        if has_drawing(p):
            return f"ALREADY_PRESENT {os.path.basename(path)}"

    fig_no = count_seq_figure(paras, anchor) + 1
    anchor_p = paras[anchor] if anchor < len(paras) else None

    # 4) paragraphe image (insere avant le prochain titre)
    if anchor_p is not None:
        img_p = anchor_p.insert_paragraph_before()
    else:
        img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.space_before = Pt(8)
    img_p.paragraph_format.space_after = Pt(2)
    keep_with_next(img_p)
    img_p.add_run().add_picture(IMG, width=Cm(WIDTH_CM))

    # 5) paragraphe legende
    if anchor_p is not None:
        cap_p = anchor_p.insert_paragraph_before()
    else:
        cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(1)
    cap_p.paragraph_format.space_after = Pt(10)
    r = cap_p.add_run("Figure ")
    r.font.name = HEAD_FONT; r.font.size = Pt(9.5); r.font.bold = True; r.font.color.rgb = PRIMARY
    _field(cap_p, 'SEQ Figure \\* ARABIC', cached=str(fig_no),
           font=HEAD_FONT, size=9.5, color=PRIMARY, bold=True)
    r2 = cap_p.add_run(f" — {CAPTION}")
    r2.font.name = HEAD_FONT; r2.font.size = Pt(9.5); r2.font.italic = True; r2.font.color.rgb = GREY

    # 6) sauvegarde + ecriture
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    bak = path.replace(".docx", f".BAK_n8n_{ts}.docx")
    shutil.copy(path, bak)
    doc.save(path)
    return f"ADDED (Figure {fig_no}) {os.path.basename(path)}  [backup: {os.path.basename(bak)}]"


if __name__ == "__main__":
    assert os.path.exists(IMG), f"image introuvable : {IMG}"
    for t in TARGETS:
        print(inject(t))
