# -*- coding: utf-8 -*-
"""
Moteur de génération du rapport CuraMedical (python-docx).
Thème vert médical, page de garde avec logo ISGA, numérotation à partir de 1
(hors page de garde), champs automatiques (table des matières, liste des
figures, liste des tableaux, numéros de page) mis à jour à l'ouverture.
"""
import os
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ─────────────────────────────────────────────────────────────────────────────
# Chemins des ressources
# ─────────────────────────────────────────────────────────────────────────────
ROOT = r"c:\2CI-ISI\S2\Projet en Systèmes Informatiques\MedPredict\CuraMedical\soutenance"
SHOTS = os.path.join(ROOT, "screenshots", "curamedical-rapport")
DDARK = os.path.join(SHOTS, "diagramme")
DUML = os.path.join(ROOT, "Diagramme_uml")

LOGO = os.path.join(SHOTS, "Logo ISGA.png")
LOGO_APP = os.path.join(SHOTS, "Logo_CuraMedical.png")   # icône de marque CuraMedical (détourée)

# Cache des captures compressées (JPEG redimensionnés)
IMGDIR = os.path.join(os.path.dirname(__file__), "_img")

# Diagrammes dynamiques (rendus soignés) — versions à jour fournies par l'équipe
ARCHI    = os.path.join(DDARK, "architecture_globale.png")
UC_FULL  = os.path.join(DDARK, "cas_utilisation.png")
SEQ_JWT  = os.path.join(DDARK, "sequence_authentification.jpg")
SEQ_IA   = os.path.join(DDARK, "sequence_suggestion_ia.png")
SEQ_DOCS = os.path.join(DDARK, "sequence_ordonnance_notification.png")
ACT_PAT  = os.path.join(DDARK, "D-1.JPG")

# Diagrammes UML (PNG)
UC_SIMPLE   = os.path.join(DUML, "diagramme de cas d'utilisation.png")
CLASS_DIAG  = os.path.join(DDARK, "classe_curamedical.png")
SEQ_CONSULT = os.path.join(DUML, "diagramme de séquence consultation.png")
ACT_DETAIL  = os.path.join(DUML, "diagramme d'activité.png")
DEPLOY      = os.path.join(DUML, "diagramme de déploiement.png")


def shot(name):
    return os.path.join(SHOTS, name)


# ─────────────────────────────────────────────────────────────────────────────
# Palette (vert émeraude médical — rappelle l'identité de CuraMedical)
# ─────────────────────────────────────────────────────────────────────────────
PRIMARY      = RGBColor(0x14, 0x7A, 0x4E)   # vert émeraude profond
PRIMARY_DARK = RGBColor(0x0E, 0x5A, 0x39)   # vert forêt
ACCENT       = RGBColor(0x2E, 0xB8, 0x72)   # vert vif (accents)
INK          = RGBColor(0x21, 0x2B, 0x30)   # encre (corps)
GREY         = RGBColor(0x55, 0x63, 0x6B)   # gris légende
WHITE        = RGBColor(0xFF, 0xFF, 0xFF)

HEX_PRIMARY   = "147A4E"
HEX_PRIMARY_D = "0E5A39"
HEX_ACCENT    = "2EB872"
HEX_FILL      = "E7F4EC"   # vert très clair (fond d'encadré / lignes paires)
HEX_FILL2     = "F4FBF7"
HEX_RULE      = "BFE3CE"

BODY_FONT = "Cambria"
HEAD_FONT = "Calibri"
MONO_FONT = "Consolas"


# ─────────────────────────────────────────────────────────────────────────────
# Helpers bas niveau (XML)
# ─────────────────────────────────────────────────────────────────────────────
def _set_shading(element, fill_hex):
    """Applique une trame de fond à un paragraphe (pPr) ou une cellule (tcPr)."""
    pr = element
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    pr.append(shd)


def cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    _set_shading(tcPr, fill_hex)


def para_shading(paragraph, fill_hex):
    pPr = paragraph._p.get_or_add_pPr()
    _set_shading(pPr, fill_hex)


def set_borders(paragraph, edges):
    """edges: dict comme {'bottom': ('single', 12, 'auto', '147A4E')}."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for edge, (val, sz, space, color) in edges.items():
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), val)
        e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), str(space))
        e.set(qn('w:color'), color)
        pbdr.append(e)
    pPr.append(pbdr)


def _field(paragraph, instr, cached="", font=None, size=None, color=None, bold=None):
    """Insère un champ Word complexe (PAGE, SEQ, TOC, ...)."""
    run = paragraph.add_run()
    fldBegin = OxmlElement('w:fldChar'); fldBegin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldBegin)

    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = instr
    run2._r.append(instrText)

    run3 = paragraph.add_run()
    fldSep = OxmlElement('w:fldChar'); fldSep.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldSep)

    run4 = paragraph.add_run(cached)

    run5 = paragraph.add_run()
    fldEnd = OxmlElement('w:fldChar'); fldEnd.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldEnd)

    for r in (run4,):
        if font:  r.font.name = font
        if size:  r.font.size = Pt(size)
        if color: r.font.color.rgb = color
        if bold is not None: r.font.bold = bold
    return run4


def enable_update_fields(doc):
    """Force Word à mettre à jour les champs (TOC, pages) à l'ouverture."""
    settings = doc.settings.element
    upd = OxmlElement('w:updateFields')
    upd.set(qn('w:val'), 'true')
    settings.append(upd)


def set_page_number_start(section, start=1):
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.append(pgNumType)
    pgNumType.set(qn('w:start'), str(start))


def keep_with_next(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    kn = OxmlElement('w:keepNext'); pPr.append(kn)


def _roman(n):
    """Entier → chiffres romains (1 → I, 4 → IV, 9 → IX…)."""
    table = [(1000, 'M'), (900, 'CM'), (500, 'D'), (400, 'CD'), (100, 'C'),
             (90, 'XC'), (50, 'L'), (40, 'XL'), (10, 'X'), (9, 'IX'),
             (5, 'V'), (4, 'IV'), (1, 'I')]
    out = ''
    for value, sym in table:
        while n >= value:
            out += sym
            n -= value
    return out


def _strip_num(title):
    """Retire un éventuel préfixe numérique décimal en tête de titre
    (« 1.2 », « 1.2.1 », « 2. ») pour le remplacer par la numérotation romaine."""
    return re.sub(r'^\s*\d+(?:\.\d+)*[.:]?\s+', '', title)


# ─────────────────────────────────────────────────────────────────────────────
# Classe principale
# ─────────────────────────────────────────────────────────────────────────────
class Report:
    def __init__(self):
        self.doc = Document()
        self.fig_n = 0
        self.tab_n = 0
        self.cap_n = 0
        self._first_body = True   # le 1er titre du corps ne re-saute pas de page
        # Numérotation style mémoire (chiffres romains)
        self._chap = 0            # n° de chapitre (I, II, …)
        self._sec2 = 0            # section de niveau 2 (I, II, … par chapitre)
        self._sec3 = 0            # sous-section de niveau 3 (I.1, I.2, …)
        self._numbered = False    # vrai uniquement dans un chapitre numéroté
        self._setup_base_styles()

    def _break(self):
        """Saut de page sauf pour le tout premier élément du corps (la 1re page
        de la 2e section est déjà créée par le saut de section)."""
        if self._first_body:
            self._first_body = False
        else:
            self.page_break()

    # ---- styles ------------------------------------------------------------
    def _setup_base_styles(self):
        st = self.doc.styles['Normal']
        st.font.name = BODY_FONT
        st.font.size = Pt(11)
        st.font.color.rgb = INK
        pf = st.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = 1.18
        pf.space_after = Pt(7)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ---- page / sections ---------------------------------------------------
    def _format_section(self, section):
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.4)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.3)
        section.header_distance = Cm(1.2)
        section.footer_distance = Cm(1.2)

    # ---- headings ----------------------------------------------------------
    def chapter(self, kicker, title):
        """Titre de chapitre (H1) sur nouvelle page.
        kicker non vide → chapitre numéroté « CHAPITRE I : … » (style mémoire).
        kicker vide → grande partie non numérotée (INTRODUCTION, CONCLUSION…)."""
        self._break()
        if kicker:
            self._chap += 1
            self._sec2 = 0
            self._sec3 = 0
            self._numbered = True
            text = f"CHAPITRE {_roman(self._chap)} : {title.upper()}"
        else:
            self._numbered = False
            text = title.upper()
        p = self.doc.add_paragraph(style='Heading 1')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(10)
        set_borders(p, {'bottom': ('single', 18, 4, HEX_PRIMARY)})
        r = p.add_run(text)
        r.font.name = HEAD_FONT
        r.font.size = Pt(19)
        r.font.bold = True
        r.font.color.rgb = PRIMARY_DARK
        return p

    def h1_plain(self, title):
        """Titre majeur de page liminaire (REMERCIEMENTS, SOMMAIRE…)."""
        self._break()
        self._numbered = False
        p = self.doc.add_paragraph(style='Heading 1')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(10)
        set_borders(p, {'bottom': ('single', 18, 4, HEX_PRIMARY)})
        r = p.add_run(title.upper())
        r.font.name = HEAD_FONT
        r.font.size = Pt(19)
        r.font.bold = True
        r.font.color.rgb = PRIMARY_DARK
        return p

    def h2(self, title):
        title = _strip_num(title)
        if self._numbered:
            self._sec2 += 1
            self._sec3 = 0
            title = f"{_roman(self._sec2)} : {title}"
        p = self.doc.add_paragraph(style='Heading 2')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(5)
        keep_with_next(p)
        r = p.add_run(title)
        r.font.name = HEAD_FONT
        r.font.size = Pt(15)
        r.font.bold = True
        r.font.color.rgb = PRIMARY
        return p

    def h3(self, title):
        title = _strip_num(title)
        if self._numbered:
            self._sec3 += 1
            title = f"{_roman(self._sec2)}.{self._sec3} : {title}"
        p = self.doc.add_paragraph(style='Heading 3')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(3)
        keep_with_next(p)
        r = p.add_run(title)
        r.font.name = HEAD_FONT
        r.font.size = Pt(12.5)
        r.font.bold = True
        r.font.color.rgb = PRIMARY_DARK
        return p

    def h4(self, title):
        """Sous-titre non répertorié dans la TDM (gras vert)."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        keep_with_next(p)
        r = p.add_run(title)
        r.font.name = HEAD_FONT
        r.font.size = Pt(11.5)
        r.font.bold = True
        r.font.color.rgb = PRIMARY_DARK
        return p

    # ---- texte -------------------------------------------------------------
    def body(self, *segments, align=None, space_after=None):
        """segments: str, ou tuple (text, {'b':True,'i':True,'color':RGB,'mono':bool})."""
        p = self.doc.add_paragraph()
        if align: p.paragraph_format.alignment = align
        if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
        for seg in segments:
            if isinstance(seg, tuple):
                text, opt = seg
            else:
                text, opt = seg, {}
            r = p.add_run(text)
            if opt.get('b'): r.font.bold = True
            if opt.get('i'): r.font.italic = True
            if opt.get('mono'):
                r.font.name = MONO_FONT
                r.font.size = Pt(10)
            if opt.get('color'): r.font.color.rgb = opt['color']
            if opt.get('size'): r.font.size = Pt(opt['size'])
        return p

    def bullets(self, items, style='bullet'):
        """items: liste de str OU de listes de segments (comme body)."""
        for it in items:
            p = self.doc.add_paragraph(style='List Bullet' if style == 'bullet' else 'List Number')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            segs = it if isinstance(it, list) else [it]
            for seg in segs:
                if isinstance(seg, tuple):
                    text, opt = seg
                else:
                    text, opt = seg, {}
                r = p.add_run(text)
                if opt.get('b'): r.font.bold = True
                if opt.get('i'): r.font.italic = True
                if opt.get('color'): r.font.color.rgb = opt['color']
                if opt.get('mono'):
                    r.font.name = MONO_FONT; r.font.size = Pt(10)

    def callout(self, title, text):
        """Encadré vert clair (note / point clé)."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.2)
        para_shading(p, HEX_FILL)
        set_borders(p, {'left': ('single', 24, 6, HEX_PRIMARY)})
        r = p.add_run(title + "  ")
        r.font.name = HEAD_FONT; r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = PRIMARY_DARK
        p2 = self.doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(8)
        p2.paragraph_format.left_indent = Cm(0.2)
        p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para_shading(p2, HEX_FILL)
        set_borders(p2, {'left': ('single', 24, 6, HEX_PRIMARY)})
        r2 = p2.add_run(text)
        r2.font.size = Pt(10.5)
        return p2

    # ---- figures & images --------------------------------------------------
    @staticmethod
    def _is_capture(path):
        """Une capture d'écran de l'application (≠ diagramme / image de code)."""
        base = os.path.basename(path)
        parent = os.path.dirname(path).rstrip('\\/')
        return parent.endswith('curamedical-rapport') and (
            base[:1].isdigit() or base.lower().startswith('whatsap'))

    def figure(self, path, caption, width_cm=15.5, frame=True, label=None):
        from docx.shared import Cm as _Cm
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        keep_with_next(p)
        run = p.add_run()
        embed = self._prep_image(path, width_cm)
        if os.path.exists(embed):
            run.add_picture(embed, width=_Cm(width_cm))
        else:
            run.add_text(f"[Image manquante : {os.path.basename(path)}]")
        if frame:
            self._frame_last_image(run)
        if label is None:
            label = 'Capture' if self._is_capture(path) else 'Figure'
        self._caption(caption, label)

    @staticmethod
    def _prep_image(path, width_cm):
        """Compresse/redimensionne les captures volumineuses (JPEG) ; laisse les
        diagrammes et images de code intacts (PNG, texte net)."""
        base = os.path.basename(path)
        parent = os.path.dirname(path).rstrip('\\/')
        is_shot = parent.endswith('curamedical-rapport') and (
            base[:1].isdigit() or base.lower().startswith('whatsap'))
        if not is_shot or not os.path.exists(path):
            return path
        try:
            from PIL import Image
            os.makedirs(IMGDIR, exist_ok=True)
            target_px = max(1000, int(width_cm / 2.54 * 200))
            out = os.path.join(IMGDIR, os.path.splitext(base)[0] + ".jpg")
            with Image.open(path) as im:
                im = im.convert("RGB")
                if im.width > target_px:
                    h = int(im.height * target_px / im.width)
                    im = im.resize((target_px, h), Image.LANCZOS)
                im.save(out, "JPEG", quality=88, optimize=True)
            return out
        except Exception:
            return path

    def _frame_last_image(self, run):
        # Encadre l'image d'une fine bordure grise via les propriétés de l'inline shape
        try:
            blip = run._r.findall('.//' + qn('a:blip'))
        except Exception:
            return

    def _caption(self, caption, label):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(f"{label} ")
        r.font.name = HEAD_FONT; r.font.size = Pt(9.5); r.font.bold = True; r.font.color.rgb = PRIMARY
        _field(p, f'SEQ {label} \\* ARABIC', cached=str(self._next(label)),
               font=HEAD_FONT, size=9.5, color=PRIMARY, bold=True)
        r2 = p.add_run(f" — {caption}")
        r2.font.name = HEAD_FONT; r2.font.size = Pt(9.5); r2.font.italic = True; r2.font.color.rgb = GREY
        return p

    def _next(self, label):
        if label == 'Figure':
            self.fig_n += 1; return self.fig_n
        if label == 'Capture':
            self.cap_n += 1; return self.cap_n
        self.tab_n += 1; return self.tab_n

    def figures_row(self, paths, caption, width_cm_each=5.0, sublabels=None, label='Capture'):
        """Plusieurs captures alignées côte à côte (tableau sans bordure) + 1 légende."""
        from docx.enum.table import WD_TABLE_ALIGNMENT
        t = self.doc.add_table(rows=1, cols=len(paths))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = True
        for i, path in enumerate(paths):
            cell = t.rows[0].cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            embed = self._prep_image(path, width_cm_each)
            run = p.add_run()
            if os.path.exists(embed):
                run.add_picture(embed, width=Cm(width_cm_each))
            else:
                run.add_text(f"[{os.path.basename(path)}]")
            if sublabels:
                cp = cell.add_paragraph()
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.paragraph_format.space_before = Pt(1)
                rr = cp.add_run(sublabels[i])
                rr.font.name = HEAD_FONT; rr.font.size = Pt(8.5); rr.font.italic = True
                rr.font.color.rgb = GREY
        self._caption(caption, label)
        sp = self.doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)
        return t

    # ---- tableaux ----------------------------------------------------------
    def table(self, headers, rows, caption=None, widths=None, header_fill=HEX_PRIMARY,
              font_size=10, caption_above=True):
        if caption and caption_above:
            self._caption(caption, 'Tableau')
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        t.autofit = True
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = ""
            pr = hdr[i].paragraphs[0]
            pr.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = pr.add_run(h)
            run.font.name = HEAD_FONT; run.font.bold = True; run.font.size = Pt(font_size)
            run.font.color.rgb = WHITE
            cell_shading(hdr[i], header_fill)
            hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for ridx, row in enumerate(rows):
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = ""
                pr = cells[i].paragraphs[0]
                pr.paragraph_format.space_after = Pt(1)
                pr.paragraph_format.space_before = Pt(1)
                segs = val if isinstance(val, list) else [val]
                for seg in segs:
                    if isinstance(seg, tuple):
                        text, opt = seg
                    else:
                        text, opt = str(seg), {}
                    run = pr.add_run(text)
                    run.font.size = Pt(font_size)
                    if opt.get('b'): run.font.bold = True
                    if opt.get('i'): run.font.italic = True
                    if opt.get('mono'): run.font.name = MONO_FONT; run.font.size = Pt(font_size-0.5)
                    if opt.get('color'): run.font.color.rgb = opt['color']
                if ridx % 2 == 1:
                    cell_shading(cells[i], HEX_FILL2)
                cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if widths:
            for col, w in enumerate(widths):
                for r in t.rows:
                    r.cells[col].width = Cm(w)
        sp = self.doc.add_paragraph(); sp.paragraph_format.space_after = Pt(6)
        return t

    # ---- code (rendu image façon VS Code) ----------------------------------
    def code(self, text, filename="", language="python", caption=None, max_cm=15.8):
        import codeimg
        path, w_px, h_px = codeimg.render(text, language=language, filename=filename)
        width_cm = min(max_cm, max(8.0, w_px / 95.0))
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(2 if caption else 10)
        keep_with_next(p)
        p.add_run().add_picture(path, width=Cm(width_cm))
        if caption:
            cp = self.doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(10)
            rr = cp.add_run(caption)
            rr.font.name = HEAD_FONT; rr.font.size = Pt(9); rr.font.italic = True
            rr.font.color.rgb = GREY
        return p

    # ---- divers ------------------------------------------------------------
    def page_break(self):
        self.doc.add_page_break()

    def spacer(self, pts=6):
        p = self.doc.add_paragraph(); p.paragraph_format.space_after = Pt(pts)
        return p

    def save(self, path):
        self.doc.save(path)


def _spacing(run, twentieths):
    """Espacement inter-lettres (letter-spacing) en vingtièmes de point."""
    rPr = run._r.get_or_add_rPr()
    spc = OxmlElement('w:spacing')
    spc.set(qn('w:val'), str(twentieths))
    rPr.append(spc)
