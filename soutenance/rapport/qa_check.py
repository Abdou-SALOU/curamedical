# -*- coding: utf-8 -*-
"""Controle qualite final du rapport (lecture seule, python-docx)."""
import os, re
from docx import Document
from docx.oxml.ns import qn

BASE = r"c:\2CI-ISI\S2\Projet en Systèmes Informatiques\MedPredict\CuraMedical\soutenance\screenshots\curamedical-rapport"
OUT  = os.path.join(BASE, "Rapport_CuraMedical.docx")
doc = Document(OUT)

# ── contenu de la section 1 (page de garde) ──
print("=== SECTION 1 (page de garde) — paragraphes non vides ===")
for p in doc.paragraphs:
    has_sect = (p._p.pPr is not None and p._p.pPr.find(qn('w:sectPr')) is not None)
    t = p.text.strip()
    has_img = bool(p._p.findall('.//' + qn('w:drawing')))
    if t or has_img:
        print(f"   {'[IMG] ' if has_img else ''}{t[:80]}")
    if has_sect:
        break

# ── stats globales ──
full = "\n".join(p.text for p in doc.paragraphs)
for tbl in doc.tables:
    for row in tbl.rows:
        for c in row.cells:
            full += "\n" + c.text
n_draw = len(doc.element.body.findall('.//' + qn('w:drawing')))
print("\n=== STATS ===")
print("   paragraphes :", len(doc.paragraphs))
print("   tableaux    :", len(doc.tables))
print("   images (drawings, corps) :", n_draw)
print("   sections    :", len(doc.sections))

# ── faits cles attendus ──
print("\n=== PRESENCE DES FAITS CLES ===")
facts = {
    "Abdou SALOU ABDOU": "Abdou SALOU ABDOU",
    "Kamara MACIRE": "Kamara MACIRE",
    "Nouridine SAWADOGO": "Nouridine SAWADOGO",
    "Dr. Soumia CHOKRI": "Soumia CHOKRI",
    "CuraMedical": "CuraMedical",
    "2025-2026 / 2025–2026": "2025",
    "Juin 2026 (soutenance)": "Juin 2026",
}
for label, needle in facts.items():
    print(f"   {'OK ' if needle in full else 'ABSENT '} {label}")

# ── drapeaux rouges ──
print("\n=== DRAPEAUX ROUGES (devrait etre 0) ===")
flags = {
    "Texte indice de champ non actualise (F9)": r"s[eé]lectionner puis F9",
    "Image manquante": r"\[Image manquante",
    "Marqueur TODO/XXX": r"\b(TODO|FIXME|XXX|TBD)\b",
    "Lorem ipsum": r"[Ll]orem ipsum",
    "Double espace suspect": r"  +[A-Za-zÀ-ÿ]",
    "Accent casse (mojibake Ã/Â)": r"[ÃÂ][\x80-\xBF]",
}
for label, pat in flags.items():
    hits = re.findall(pat, full)
    print(f"   {'OK (0)' if not hits else 'A VOIR (%d)' % len(hits)}  {label}" + (f"  ex: {hits[:3]}" if hits else ""))
